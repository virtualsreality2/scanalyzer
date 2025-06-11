"""DOCX parser for Word document security reports."""

import io
import re
import logging
from typing import AsyncIterator, Optional, Dict, Any, List
from pathlib import Path

from app.parsers.base import AbstractParser, ParserMetadata, ParserCapabilities
from app.parsers.registry import register_parser
from app.models.finding import Finding, SeverityLevel
from app.core.exceptions import ParseError
from .patterns import PatternMatcher, SeverityMapper, ConfidenceScorer, TABLE_HEADER_MAPPINGS

logger = logging.getLogger(__name__)

# Import DOCX library with fallback
try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed - DOCX parsing disabled")


@register_parser
class DocxParser(AbstractParser):
    """Parser for DOCX (Word) security reports."""
    
    def __init__(self):
        self.pattern_matcher = PatternMatcher()
        self.severity_mapper = SeverityMapper()
        self.confidence_scorer = ConfidenceScorer()
    
    @classmethod
    def get_metadata(cls) -> ParserMetadata:
        return ParserMetadata(
            name="DOCX Security Report Parser",
            tool="docx",
            supported_versions=["*"],
            file_extensions=[".docx"],
            confidence_threshold=0.7
        )
    
    @classmethod
    def get_capabilities(cls) -> ParserCapabilities:
        return ParserCapabilities(
            streaming=True,
            format_detection=True,
            auto_remediation=False,
            compliance_mapping=False
        )
    
    async def can_parse(self, file_path: Path, content_preview: bytes) -> float:
        """Check if file is a DOCX."""
        # DOCX files are actually ZIP archives
        # Check for ZIP magic bytes
        if content_preview.startswith(b'PK\x03\x04'):
            # Could be DOCX, check file extension
            if file_path.suffix.lower() == '.docx':
                return 0.95
            return 0.3  # Could be any ZIP file
        
        # Check file extension only
        if file_path.suffix.lower() == '.docx':
            return 0.7
        
        return 0.0
    
    async def parse_stream(self, file_stream: io.IOBase) -> AsyncIterator[Finding]:
        """Parse DOCX security report."""
        if not HAS_DOCX:
            raise ParseError("python-docx not installed - cannot parse DOCX files")
        
        try:
            # Load document
            doc = Document(file_stream)
            
            # Extract all text and findings
            findings = []
            
            # Process paragraphs
            full_text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n"
            
            # Extract findings from text
            if full_text:
                text_findings = self.pattern_matcher.extract_findings(full_text)
                findings.extend(text_findings)
            
            # Process tables
            table_findings = await self._extract_from_tables(doc.tables)
            findings.extend(table_findings)
            
            # Process structured sections
            section_findings = await self._extract_from_sections(doc)
            findings.extend(section_findings)
            
            # Convert findings to Finding objects
            for finding_data in findings:
                confidence = self.confidence_scorer.calculate_confidence(finding_data)
                
                # Only yield findings above confidence threshold
                if confidence >= 0.5:
                    yield await self._create_finding(finding_data, confidence)
        
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise ParseError(f"Failed to parse DOCX: {str(e)}")
    
    async def _extract_from_tables(self, tables: List[Table]) -> List[Dict[str, Any]]:
        """Extract findings from DOCX tables."""
        findings = []
        
        for table in tables:
            if not table.rows:
                continue
            
            # Extract headers from first row
            headers = []
            for cell in table.rows[0].cells:
                headers.append(cell.text.lower().strip())
            
            # Map headers to finding fields
            field_mapping = self._map_table_headers(headers)
            if not field_mapping:
                # Try to extract as key-value pairs
                kv_findings = await self._extract_key_value_table(table)
                findings.extend(kv_findings)
                continue
            
            # Extract findings from data rows
            for row in table.rows[1:]:
                finding = {}
                cells = row.cells
                
                for i, cell in enumerate(cells):
                    if i < len(headers) and headers[i] in field_mapping:
                        field = field_mapping[headers[i]]
                        value = cell.text.strip()
                        if value:
                            finding[field] = value
                
                if finding.get('title') or finding.get('description'):
                    # Map severity
                    if 'severity' in finding:
                        finding['severity'] = self.severity_mapper.map_severity(
                            finding['severity']
                        ).value
                    
                    findings.append(finding)
        
        return findings
    
    async def _extract_key_value_table(self, table: Table) -> List[Dict[str, Any]]:
        """Extract findings from key-value style tables."""
        findings = []
        current_finding = {}
        
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.lower().strip()
                value = row.cells[1].text.strip()
                
                if not key or not value:
                    continue
                
                # Check if this is a new finding
                if any(term in key for term in ['finding', 'issue', 'vulnerability', 'title']):
                    if current_finding:
                        findings.append(current_finding)
                    current_finding = {'title': value}
                
                # Map common fields
                elif any(term in key for term in ['severity', 'risk', 'priority']):
                    current_finding['severity'] = self.severity_mapper.map_severity(value).value
                
                elif any(term in key for term in ['description', 'details', 'summary']):
                    current_finding['description'] = value
                
                elif any(term in key for term in ['recommendation', 'remediation', 'fix']):
                    current_finding['recommendation'] = value
                
                elif any(term in key for term in ['file', 'resource', 'location']):
                    current_finding['resource'] = value
        
        # Add last finding
        if current_finding and (current_finding.get('title') or current_finding.get('description')):
            findings.append(current_finding)
        
        return findings
    
    async def _extract_from_sections(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract findings from document sections."""
        findings = []
        current_section = None
        current_finding = {}
        
        for element in self._iter_block_items(doc):
            if isinstance(element, Paragraph):
                text = element.text.strip()
                if not text:
                    continue
                
                # Check for section headers
                if element.style and element.style.name:
                    style_name = element.style.name.lower()
                    
                    # Major section header - might be a new finding
                    if 'heading' in style_name and any(term in text.lower() for term in ['finding', 'issue', 'vulnerability']):
                        if current_finding:
                            findings.append(current_finding)
                        current_finding = {'title': text}
                        current_section = 'finding'
                    
                    # Subsection headers
                    elif current_section == 'finding':
                        if any(term in text.lower() for term in ['severity', 'risk', 'priority']):
                            current_section = 'severity'
                        elif any(term in text.lower() for term in ['description', 'details']):
                            current_section = 'description'
                        elif any(term in text.lower() for term in ['recommendation', 'remediation']):
                            current_section = 'recommendation'
                
                # Regular paragraph - add to current section
                elif current_section:
                    if current_section == 'severity':
                        current_finding['severity'] = self.severity_mapper.map_severity(text).value
                    elif current_section == 'description':
                        current_finding['description'] = current_finding.get('description', '') + '\n' + text
                    elif current_section == 'recommendation':
                        current_finding['recommendation'] = current_finding.get('recommendation', '') + '\n' + text
        
        # Add last finding
        if current_finding and (current_finding.get('title') or current_finding.get('description')):
            findings.append(current_finding)
        
        return findings
    
    def _iter_block_items(self, parent):
        """Iterate through block items (paragraphs and tables) in document order."""
        from docx.document import Document as DocumentType
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        
        if isinstance(parent, DocumentType):
            parent_elm = parent.element.body
        else:
            raise ValueError("Unknown parent type")
        
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)
    
    def _map_table_headers(self, headers: List[str]) -> Dict[str, str]:
        """Map table headers to finding fields."""
        mapping = {}
        
        for field, variations in TABLE_HEADER_MAPPINGS.items():
            for header in headers:
                header_clean = header.lower().strip()
                for variation in variations:
                    if variation in header_clean or header_clean in variation:
                        mapping[header] = field
                        break
                if header in mapping:
                    break
        
        return mapping
    
    async def _create_finding(self, finding_data: Dict[str, Any], confidence: float) -> Finding:
        """Create Finding object from extracted data."""
        # Extract severity
        severity_text = finding_data.get('severity', 'MEDIUM')
        if severity_text in ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW']:
            severity = SeverityLevel[severity_text]
        else:
            severity = self.severity_mapper.map_severity(severity_text)
        
        # Build description
        description = finding_data.get('description', '').strip()
        if finding_data.get('recommendation'):
            recommendation = finding_data['recommendation'].strip()
            if recommendation:
                description += f"\n\nRecommendation: {recommendation}"
        
        # Extract resource
        resource = finding_data.get('resource', 'Document')
        
        # Build metadata
        metadata = {
            'confidence': confidence,
            'source': 'docx_parser',
            'extraction_method': 'structured_extraction'
        }
        
        # Add references if available
        if finding_data.get('references'):
            metadata['references'] = finding_data['references']
        
        # Add vulnerability type if identified
        if finding_data.get('vulnerability_type'):
            metadata['vulnerability_type'] = finding_data['vulnerability_type']
        
        return Finding(
            title=finding_data.get('title', 'Security Finding'),
            severity=severity,
            description=description,
            resource_name=resource,
            metadata=metadata
        )